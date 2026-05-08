"""
基础格式检查 (Check 1-5) —— 检测文档排版问题，不涉及内容幻觉。
"""

from __future__ import annotations

import re
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from docx.document import Document

from ..engine.config import GEHDConfig


def check_markdown(doc: Document) -> list[str]:
    """Check 1: 检测 Markdown 语法符号残留。"""
    patterns = [r'\*\*', r'~~', r'`{1,3}', r'^#{1,6}\s']
    issues: list[str] = []

    for i, para in enumerate(doc.paragraphs):
        t = para.text
        for pat in patterns:
            if re.search(pat, t):
                issues.append(f'[MARKDOWN] P{i + 1} 含markdown({pat}): "{t[:50]}"')

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pat in patterns:
                    if re.search(pat, cell.text):
                        issues.append(f'[MARKDOWN] T{ti + 1}[{ri + 1},{ci + 1}] markdown残留')

    return issues


def check_empty_table_rows(doc: Document) -> list[str]:
    """Check 2: 检测空表格行。"""
    issues: list[str] = []

    for ti, table in enumerate(doc.tables):
        empty_rows = sum(
            1 for row in table.rows if all(cell.text.strip() == '' for cell in row.cells)
        )
        if empty_rows > 0:
            issues.append(f'[空表行] 表格{ti + 1}: {empty_rows}个空行 / 共{len(table.rows)}行')

    return issues


def check_blank_paragraphs(doc: Document, config: GEHDConfig) -> list[str]:
    """Check 3: 检测大面积连续空白段落。"""
    issues: list[str] = []
    consecutive = 0
    max_consec = 0
    empty_positions: list[int] = []

    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == '':
            consecutive += 1
            max_consec = max(max_consec, consecutive)
            empty_positions.append(i + 1)
        else:
            consecutive = 0

    limit = config.max_consecutive_blank_paragraphs
    if max_consec > limit:
        issues.append(
            f'[空白] 连续{max_consec}个空段 (P{min(empty_positions)}-P{max(empty_positions)})'
        )

    cover_empties = sum(1 for p in list(doc.paragraphs)[:8] if p.text.strip() == '')
    if cover_empties > limit:
        issues.append(f'[封面空白] 前8段中{cover_empties}个空段(建议<={limit})')

    return issues


def check_emoji(doc: Document) -> list[str]:
    """Check 4: 检测 Emoji 和方块字符。"""
    emoji_ranges = [
        (0x1F300, 0x1F9FF),
        (0x2600, 0x26FF),
        (0x2700, 0x27BF),
        (0xFE00, 0xFE0F),
        (0x1FA00, 0x1FA6F),
        (0x1FA70, 0x1FAFF),
        (0x2300, 0x23FF),
    ]
    issues: list[str] = []

    for i, para in enumerate(doc.paragraphs):
        for ch in para.text:
            cp = ord(ch)
            if any(s <= cp <= e for s, e in emoji_ranges):
                issues.append(f'[Emoji] P{i + 1} U+{cp:04X}: "{para.text[:45]}"')
                break

    return issues


def check_long_text(doc: Document, config: GEHDConfig) -> list[str]:
    """Check 5: 检测超长文本（表格单元格截断预警）。"""
    warnings: list[str] = []

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if len(cell.text) > config.long_text_threshold_chars:
                    warnings.append(
                        f'[长文本] T{ti + 1}[{ri + 1},{ci + 1}] {len(cell.text)}字可能截断'
                    )

    return warnings
